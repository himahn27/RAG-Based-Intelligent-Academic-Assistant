import os
import pandas as pd
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from config import CHUNK_SIZE, CHUNK_OVERLAP


# ----------------------------
# EXTRACT TEXT FROM FILE
# ----------------------------
def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    # ---------------- PDF ----------------
    if ext == ".pdf":
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text

    # ---------------- DOCX ----------------
    elif ext == ".docx":
        doc = DocxDocument(file_path)
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    # ---------------- TXT ----------------
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    # ---------------- CSV ----------------
    elif ext == ".csv":
        df = pd.read_csv(file_path)
        return df.to_string(index=False)

    # ---------------- EXCEL (MULTI-SHEET SUPPORT) ----------------
    elif ext in [".xlsx", ".xls"]:
        sheets = pd.read_excel(file_path, sheet_name=None)  # Load all sheets
        text = ""
        for sheet_name, df in sheets.items():
            text += f"\nSheet: {sheet_name}\n"

            for _,row in df.iterrows():
                row_text = " ,".join([f"{col}: {row[col]}" for col in df.columns])
                text += row_text + "\n"

            text +="\n"    
        return text

    else:
        raise ValueError(f"Unsupported file format: {ext}")


# ----------------------------
# LOAD AND SPLIT INTO DOCUMENT OBJECTS
# ----------------------------
def load_and_split(file_path):
    text = extract_text(file_path)

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP
    )

    chunks = splitter.split_text(text)

    # Convert chunks to LangChain Document objects with metadata
    documents = [
        Document(
            page_content=chunk,
            metadata={
                "source": os.path.basename(file_path),
                "file_type": os.path.splitext(file_path)[1].lower()
            }
        )
        for chunk in chunks
    ]

    return documents